"""Resume text extraction and structured field parsing.

Supports PDF, DOCX and TXT. Extraction is pure-Python (pypdf / python-docx)
so it runs anywhere without native build tools.
"""

from __future__ import annotations

import io
import re
from typing import Dict, List

from pypdf import PdfReader
from docx import Document

from .skills import SKILL_ALIASES


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #
def extract_text(filename: str, data: bytes) -> str:
    """Extract raw text from an uploaded resume file."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError("Unsupported file type. Use PDF, DOCX or TXT.")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# Field parsing
# --------------------------------------------------------------------------- #
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(
    r"(?:(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{3,5}\)?[\s-]?)?\d{3}[\s-]?\d{3,4})"
)
URL_RE = re.compile(r"(?:https?://|www\.)[^\s)]+", re.IGNORECASE)
LINKEDIN_RE = re.compile(r"(?:linkedin\.com/in/[^\s)]+)", re.IGNORECASE)
GITHUB_RE = re.compile(r"(?:github\.com/[^\s)]+)", re.IGNORECASE)
# Fallback for icon-embedded handles, e.g. "/linkedinlakshya-agrawal" or
# "/githubLuckyMan22-SuperMan" produced when a PDF hides the full URL behind
# an icon glyph. Handle is a run of URL-safe chars until whitespace/next glyph.
LINKEDIN_HANDLE_RE = re.compile(r"/?linkedin[/:]?([A-Za-z0-9][A-Za-z0-9._-]{1,})", re.IGNORECASE)
GITHUB_HANDLE_RE = re.compile(r"/?github[/:]?([A-Za-z0-9][A-Za-z0-9._-]{1,})", re.IGNORECASE)

DEGREE_KEYWORDS = [
    "b.tech", "btech", "b.e", "bachelor", "b.sc", "bsc", "bca",
    "m.tech", "mtech", "master", "m.sc", "msc", "mca", "mba",
    "ph.d", "phd", "doctorate", "diploma", "12th", "10th", "intermediate",
]

SECTION_HEADERS = {
    "experience": ["experience", "work experience", "employment", "professional experience", "internship", "internships"],
    "education": ["education", "academic", "qualification", "academics"],
    "projects": ["projects", "project work", "personal projects"],
    "skills": ["skills", "technical skills", "technologies", "tech stack"],
}


NAME_WORD_RE = re.compile(r"^[A-Z][a-zA-Z'.\-]+$")


def _leading_name(line: str) -> str | None:
    """Return the leading person-name portion of a line, or None.

    Resumes often concatenate the name with contact icons/info on one line,
    e.g. "Lakshya Agrawal/envelOpelakshyaus4@gmail.com ...". We slice off
    everything from the first contact/icon delimiter, then keep the leading
    run of capitalised name-like words.
    """
    # Cut at the first delimiter that signals contact info / icons start.
    cut = len(line)
    for i, ch in enumerate(line):
        if ch in "@/|,:" or ord(ch) > 127 or ch.isdigit():
            cut = i
            break
    candidate = line[:cut].strip()
    if not candidate:
        return None
    words = candidate.split()
    name_words = []
    for w in words:
        if NAME_WORD_RE.match(w) or (w.isupper() and w.isalpha()):
            name_words.append(w.title() if w.isupper() else w)
        else:
            break
    if 1 < len(name_words) <= 4:
        joined = " ".join(name_words)
        if joined.lower() not in ("curriculum vitae", "resume"):
            return joined
    return None


def _guess_name(text: str, email: str | None) -> str:
    """Heuristic: first line whose leading tokens look like a person's name."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for ln in lines[:8]:
        name = _leading_name(ln)
        if name:
            return name
    # Fallback: derive from email local part.
    if email:
        local = email.split("@")[0]
        local = re.sub(r"[._0-9]+", " ", local).strip()
        if local:
            return local.title()
    return "Unknown"


def _extract_skills(text: str) -> List[str]:
    lower = " " + text.lower() + " "
    found = []
    for canonical, aliases in SKILL_ALIASES.items():
        for alias in aliases:
            pattern = r"(?<![a-zA-Z0-9+#.])" + re.escape(alias) + r"(?![a-zA-Z0-9+#])"
            if re.search(pattern, lower):
                found.append(canonical)
                break
    return sorted(set(found))


def _extract_education(text: str) -> List[str]:
    results = []
    for ln in text.splitlines():
        low = ln.lower()
        if any(kw in low for kw in DEGREE_KEYWORDS):
            clean = ln.strip()
            if clean and clean not in results:
                results.append(clean)
    return results[:6]


def _extract_section(text: str, keys: List[str]) -> str:
    """Return the text block under a section header until the next header."""
    lines = text.splitlines()
    all_headers = [h for hs in SECTION_HEADERS.values() for h in hs]
    start = None
    for i, ln in enumerate(lines):
        low = ln.strip().lower().rstrip(":")
        if low in keys:
            start = i + 1
            break
    if start is None:
        return ""
    collected = []
    for ln in lines[start:]:
        low = ln.strip().lower().rstrip(":")
        if low in all_headers and low not in keys:
            break
        collected.append(ln)
    return "\n".join(collected).strip()


def _extract_linkedin(text: str) -> str | None:
    m = LINKEDIN_RE.search(text)
    if m:
        return m.group(0)
    m = LINKEDIN_HANDLE_RE.search(text)
    if m:
        return f"linkedin.com/in/{m.group(1)}"
    return None


def _extract_github(text: str) -> str | None:
    m = GITHUB_RE.search(text)
    if m:
        return m.group(0)
    m = GITHUB_HANDLE_RE.search(text)
    if m:
        return f"github.com/{m.group(1)}"
    return None


def parse_resume(filename: str, data: bytes) -> Dict:
    """Full pipeline: bytes -> structured resume dict."""
    text = extract_text(filename, data)
    emails = EMAIL_RE.findall(text)
    phones = PHONE_RE.findall(text)
    # Filter phone matches to plausible lengths.
    phones = [p.strip() for p in phones if len(re.sub(r"\D", "", p)) >= 10]

    email = emails[0] if emails else None

    return {
        "name": _guess_name(text, email),
        "email": email,
        "phone": phones[0] if phones else None,
        "linkedin": _extract_linkedin(text),
        "github": _extract_github(text),
        "skills": _extract_skills(text),
        "education": _extract_education(text),
        "experience_section": _extract_section(text, SECTION_HEADERS["experience"]),
        "projects_section": _extract_section(text, SECTION_HEADERS["projects"]),
        "word_count": len(text.split()),
        "raw_text": text,
    }
